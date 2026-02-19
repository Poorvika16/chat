import os
import faiss
import numpy as np
import pytesseract
from docx import Document
from sentence_transformers import SentenceTransformer
from PIL import Image
import io

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

DOC_FOLDER = "docs"
IMAGE_FOLDER = "static/images"

os.makedirs(IMAGE_FOLDER, exist_ok=True)

print("Loading embedding model...")
model = SentenceTransformer("BAAI/bge-small-en")

metadata = []
embeddings = []
image_count = 0

doc_files = [f for f in os.listdir(DOC_FOLDER) if f.endswith(".docx")]

if not doc_files:
    raise Exception("No DOCX files found")

print("Documents found:", doc_files)

for file in doc_files:

    print("\nProcessing:", file)

    doc_path = os.path.join(DOC_FOLDER, file)
    doc = Document(doc_path)

    current_heading = "General"

    for para in doc.paragraphs:

        text = para.text.strip()

        if not text:
            continue

        if para.style.name.startswith("Heading"):
            current_heading = text
            continue

        chunk = f"{current_heading}\n{text}"

        metadata.append({
            "content": chunk,
            "type": "text",
            "image": None,
            "source_doc": file
        })

        emb = model.encode(chunk, convert_to_numpy=True)
        embeddings.append(emb)

    print("Extracting tables...")

    for table in doc.tables:

        table_rows = []

        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                table_rows.append(" | ".join(row_data))

        if table_rows:

            table_text = f"{current_heading}\n" + "\n".join(table_rows)

            metadata.append({
                "content": table_text,
                "type": "text",
                "image": None,
                "source_doc": file
            })

            emb = model.encode(table_text, convert_to_numpy=True)
            embeddings.append(emb)

    print("Extracting images...")

    for rel in doc.part.rels.values():

        if "image" in rel.target_ref:

            image_data = rel.target_part.blob
            image = Image.open(io.BytesIO(image_data))

            width, height = image.size

            if width < 200 or height < 200:
                continue

 
            try:
                ocr_text = pytesseract.image_to_string(image).strip()
            except:
                ocr_text = ""

            if len(ocr_text) < 5:
                continue

            image_name = f"doc_image_{image_count}.png"
            image_path = os.path.join(IMAGE_FOLDER, image_name)
            image.save(image_path)

            image_context = (
                f"{current_heading} diagram figure flowchart architecture illustration "
                f"{ocr_text}"
            )

            metadata.append({
                "content": image_context,
                "type": "image",
                "image": image_name,
                "source_doc": file
            })

            emb = model.encode(image_context, convert_to_numpy=True)
            embeddings.append(emb)

            image_count += 1

print("\nTotal images indexed:", image_count)

print("\nBuilding FAISS index...")

embeddings = np.array(embeddings).astype("float32")

faiss.normalize_L2(embeddings)

dimension = embeddings.shape[1]

index = faiss.IndexFlatIP(dimension)
index.add(embeddings)

faiss.write_index(index, "doc_index.faiss")
np.save("doc_metadata.npy", metadata)

print("Index built successfully")
